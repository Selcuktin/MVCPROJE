"""
Service Layer: Business logic and data processing operations.
Bu dosya ders işlemleri için business logic ve veri işleme operasyonlarını içerir.
"""
from django.db.models import Q, Count
from django.utils import timezone
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.shortcuts import get_object_or_404
import openpyxl
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import csv
from io import BytesIO, StringIO
import os
import re
import difflib
from typing import List, Dict, Any, Optional

from .models import Course, CourseGroup, Enrollment, AssignmentHistory, ExampleQuestion, Quiz, QuizQuestion, QuizChoice
from apps.students.models import Student
from apps.teachers.models import Teacher


class CourseService:
    """Course business logic service"""
    
    def get_filtered_courses(self, filters):
        """Get courses with applied filters"""
        queryset = Course.objects.filter(status='active')
        
        if filters.get('search'):
            queryset = queryset.filter(
                Q(code__icontains=filters['search']) |
                Q(name__icontains=filters['search']) |
                Q(department__icontains=filters['search'])
            )
        
        if filters.get('department'):
            queryset = queryset.filter(department__icontains=filters['department'])
            
        if filters.get('semester'):
            queryset = queryset.filter(semester=filters['semester'])
            
        return queryset.order_by('code')
    
    def get_course_with_details(self, course):
        """Get course with related enrollments and content"""
        enrollments = Enrollment.objects.filter(
            group__course=course,
            status='enrolled'
        ).select_related('student', 'group').order_by('student__first_name')
        
        # Get course contents
        try:
            from .models import CourseContent
            contents = CourseContent.objects.filter(
                course=course,
                is_active=True
            ).order_by('week_number', 'upload_date')
        except:
            contents = []
        
        return {
            'course': course,
            'enrollments': enrollments,
            'enrolled_count': enrollments.count(),
            'contents': contents,
            'teacher': course.groups.filter(status='active').first().teacher if course.groups.filter(status='active').exists() else None
        }
    
    def create_course(self, form_data):
        """Create new course"""
        course = Course.objects.create(**form_data)
        return course
    
    def update_course(self, course, form_data):
        """Update existing course"""
        for key, value in form_data.items():
            setattr(course, key, value)
        course.save()
        return course
    
    def delete_course(self, course):
        """Delete course (soft delete)"""
        course.status = 'inactive'
        course.save()
        return True


class ReportService:
    """Report generation service"""
    
    def generate_student_report(self, format_type='pdf'):
        """Generate student report in specified format"""
        enrollments = Enrollment.objects.select_related('student', 'group__course', 'group__teacher')
        
        if format_type == 'pdf':
            return self._generate_pdf_report(enrollments, 'student_report')
        elif format_type == 'excel':
            return self._generate_excel_report(enrollments, 'student_report')
        elif format_type == 'csv':
            return self._generate_csv_report(enrollments, 'student_report')
        else:
            raise ValueError("Unsupported format type")
    
    def generate_course_report(self, course, format_type='pdf'):
        """Generate course report"""
        enrollments = Enrollment.objects.filter(group__course=course).select_related('student', 'group')
        
        if format_type == 'pdf':
            return self._generate_pdf_report(enrollments, f'course_report_{course.code}')
        elif format_type == 'excel':
            return self._generate_excel_report(enrollments, f'course_report_{course.code}')
        elif format_type == 'csv':
            return self._generate_csv_report(enrollments, f'course_report_{course.code}')
    
    def _generate_pdf_report(self, data, filename):
        """Generate PDF report"""
        buffer = BytesIO()
        p = canvas.Canvas(buffer, pagesize=letter)
        
        # PDF content generation
        y = 750
        p.drawString(100, y, f"Report: {filename}")
        y -= 30
        
        for item in data:
            if hasattr(item, 'student'):
                text = f"{item.student.full_name} - {item.group.course.name}"
                p.drawString(100, y, text)
                y -= 20
                if y < 100:
                    p.showPage()
                    y = 750
        
        p.save()
        buffer.seek(0)
        
        response = HttpResponse(buffer, content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{filename}.pdf"'
        return response
    
    def _generate_excel_report(self, data, filename):
        """Generate Excel report"""
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Report"
        
        # Headers
        headers = ['Student', 'Course', 'Teacher', 'Grade', 'Status']
        for col, header in enumerate(headers, 1):
            ws.cell(row=1, column=col, value=header)
        
        # Data
        for row, item in enumerate(data, 2):
            if hasattr(item, 'student'):
                ws.cell(row=row, column=1, value=item.student.full_name)
                ws.cell(row=row, column=2, value=item.group.course.name)
                ws.cell(row=row, column=3, value=item.group.teacher.full_name)
                ws.cell(row=row, column=4, value=getattr(item, 'grade', 'N/A'))
                ws.cell(row=row, column=5, value=item.get_status_display())
        
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
        response['Content-Disposition'] = f'attachment; filename="{filename}.xlsx"'
        wb.save(response)
        return response
    
    def _generate_csv_report(self, data, filename):
        """Generate CSV report"""
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = f'attachment; filename="{filename}.csv"'
        
        writer = csv.writer(response)
        writer.writerow(['Student', 'Course', 'Teacher', 'Grade', 'Status'])
        
        for item in data:
            if hasattr(item, 'student'):
                writer.writerow([
                    item.student.full_name,
                    item.group.course.name,
                    item.group.teacher.full_name,
                    getattr(item, 'grade', 'N/A'),
                    item.get_status_display()
                ])
        
        return response


class TeacherCourseAssignmentService:
    """Teacher-Course Assignment business logic service"""
    
    def get_current_assignments(self, filters=None):
        """Get current teacher-course assignments with filters"""
        queryset = CourseGroup.objects.filter(status='active').select_related('course', 'teacher')
        
        if filters:
            if filters.get('search'):
                queryset = queryset.filter(
                    Q(course__code__icontains=filters['search']) |
                    Q(course__name__icontains=filters['search']) |
                    Q(teacher__first_name__icontains=filters['search']) |
                    Q(teacher__last_name__icontains=filters['search']) |
                    Q(teacher__department__icontains=filters['search'])
                )
            
            if filters.get('teacher_id'):
                queryset = queryset.filter(teacher_id=filters['teacher_id'])
            
            if filters.get('course_id'):
                queryset = queryset.filter(course_id=filters['course_id'])
            
            if filters.get('department'):
                queryset = queryset.filter(course__department__icontains=filters['department'])
            
            if filters.get('semester'):
                queryset = queryset.filter(semester__icontains=filters['semester'])
        
        return queryset.order_by('-semester', 'course__code')
    
    def get_assignment_history(self, filters=None, limit=50):
        """Get assignment history with filters"""
        queryset = AssignmentHistory.objects.select_related(
            'course_group__course', 'teacher', 'performed_by'
        )
        
        if filters:
            if filters.get('teacher_id'):
                queryset = queryset.filter(teacher_id=filters['teacher_id'])
            
            if filters.get('course_id'):
                queryset = queryset.filter(course_group__course_id=filters['course_id'])
            
            if filters.get('action'):
                queryset = queryset.filter(action=filters['action'])
        
        return queryset[:limit]
    
    def check_compatibility(self, teacher, course):
        """Check if teacher is compatible with course (department match)"""
        compatibility = {
            'is_compatible': False,
            'score': 0,
            'reasons': []
        }
        
        # Check department match
        if teacher.department.lower() == course.department.lower():
            compatibility['is_compatible'] = True
            compatibility['score'] = 100
            compatibility['reasons'].append('Bölüm uyumu: Mükemmel')
        elif teacher.department.lower() in course.department.lower() or course.department.lower() in teacher.department.lower():
            compatibility['is_compatible'] = True
            compatibility['score'] = 70
            compatibility['reasons'].append('Bölüm uyumu: Kısmen uyumlu')
        else:
            compatibility['reasons'].append(f'Bölüm uyumsuzluğu: Öğretmen ({teacher.department}) - Ders ({course.department})')
        
        # Check if teacher already has this course
        existing_assignments = CourseGroup.objects.filter(
            teacher=teacher,
            course=course,
            status='active'
        ).count()
        
        if existing_assignments > 0:
            compatibility['reasons'].append(f'Öğretmen bu dersi zaten veriyor ({existing_assignments} grup)')
            compatibility['score'] = max(0, compatibility['score'] - 20)
        
        return compatibility
    
    def check_schedule_conflicts(self, teacher, course, semester, schedule):
        """Check for schedule conflicts with teacher's other courses"""
        conflicts = []
        
        # Get teacher's other active course groups in the same semester
        other_groups = CourseGroup.objects.filter(
            teacher=teacher,
            semester=semester,
            status='active'
        ).exclude(course=course)
        
        for group in other_groups:
            # Simple schedule conflict check (can be enhanced with proper time parsing)
            if group.schedule and schedule:
                # Check if schedules overlap (basic check)
                if self._schedules_overlap(group.schedule, schedule):
                    conflicts.append({
                        'group': group,
                        'conflict_type': 'Zaman çakışması',
                        'message': f"{group.course.code} ({group.schedule}) ile çakışma var"
                    })
        
        return conflicts
    
    def _schedules_overlap(self, schedule1, schedule2):
        """Check if two schedules overlap (simplified)"""
        # This is a simplified check - in production, proper time parsing should be used
        days1 = self._extract_days(schedule1)
        days2 = self._extract_days(schedule2)
        
        # Check if same day
        if set(days1) & set(days2):
            # Check if time ranges overlap (simplified)
            times1 = self._extract_times(schedule1)
            times2 = self._extract_times(schedule2)
            
            if times1 and times2:
                # Simple overlap check
                return True  # Simplified - should parse times properly
        
        return False
    
    def _extract_days(self, schedule):
        """Extract days from schedule string"""
        days_map = {
            'pazartesi': 'monday', 'salı': 'tuesday', 'çarşamba': 'wednesday',
            'perşembe': 'thursday', 'cuma': 'friday', 'cumartesi': 'saturday', 'pazar': 'sunday'
        }
        days = []
        schedule_lower = schedule.lower()
        for tr_day, en_day in days_map.items():
            if tr_day in schedule_lower:
                days.append(en_day)
        return days
    
    def _extract_times(self, schedule):
        """Extract time ranges from schedule string"""
        import re
        # Match time patterns like "09:00-12:00"
        times = re.findall(r'\d{2}:\d{2}', schedule)
        return times if len(times) >= 2 else None
    
    def assign_course_to_teacher(self, course, teacher, semester, classroom, schedule, performed_by):
        """Assign course to teacher"""
        # Check compatibility
        compatibility = self.check_compatibility(teacher, course)
        
        # Check conflicts
        conflicts = self.check_schedule_conflicts(teacher, course, semester, schedule)
        
        # Create course group
        course_group = CourseGroup.objects.create(
            course=course,
            teacher=teacher,
            semester=semester,
            classroom=classroom,
            schedule=schedule,
            status='active'
        )
        
        # Log assignment history
        action = 'bulk_assign' if hasattr(self, '_bulk_mode') else 'assign'
        AssignmentHistory.objects.create(
            course_group=course_group,
            teacher=teacher,
            action=action,
            description=f"{course.code} - {course.name} dersi {teacher.full_name} öğretmenine atandı.",
            performed_by=performed_by
        )
        
        return {
            'success': True,
            'course_group': course_group,
            'compatibility': compatibility,
            'conflicts': conflicts,
            'warnings': [] if compatibility['is_compatible'] and not conflicts else ['Uyumluluk uyarıları var']
        }
    
    def remove_assignment(self, course_group_id, performed_by):
        """Remove course assignment from teacher"""
        course_group = get_object_or_404(CourseGroup, pk=course_group_id)
        teacher = course_group.teacher
        course = course_group.course
        
        # Log removal history
        AssignmentHistory.objects.create(
            course_group=course_group,
            teacher=teacher,
            action='remove',
            description=f"{course.code} - {course.name} dersi {teacher.full_name} öğretmeninden kaldırıldı.",
            performed_by=performed_by
        )
        
        # Soft delete - set status to inactive
        course_group.status = 'inactive'
        course_group.save()
        
        return {'success': True, 'message': 'Atama başarıyla kaldırıldı'}
    
    def bulk_assign(self, course_ids, teacher_ids, semester, classroom, schedule, performed_by):
        """Bulk assign courses to teachers"""
        self._bulk_mode = True
        results = []
        
        for course_id in course_ids:
            course = Course.objects.get(pk=course_id)
            for teacher_id in teacher_ids:
                teacher = Teacher.objects.get(pk=teacher_id)
                result = self.assign_course_to_teacher(
                    course, teacher, semester, classroom, schedule, performed_by
                )
                results.append(result)
        
        delattr(self, '_bulk_mode')
        return results
    
    def bulk_remove(self, course_group_ids, performed_by):
        """Bulk remove course assignments"""
        results = []
        for group_id in course_group_ids:
            result = self.remove_assignment(group_id, performed_by)
            results.append(result)
        return results
    
    def get_teacher_availability(self, teacher):
        """Get teacher's availability status"""
        active_groups = CourseGroup.objects.filter(teacher=teacher, status='active')
        
        total_courses = active_groups.values('course').distinct().count()
        total_groups = active_groups.count()
        total_students = Enrollment.objects.filter(
            group__in=active_groups,
            status='enrolled'
        ).count()
        
        return {
            'teacher': teacher,
            'total_courses': total_courses,
            'total_groups': total_groups,
            'total_students': total_students,
            'is_available': total_groups < 10,  # Max 10 groups
            'workload_percentage': min(100, (total_groups / 10) * 100)
        }


def solve_question_with_ai(question: ExampleQuestion) -> str:
    """Gemini entegrasyonu: ENV'den GEMINI_API_KEY varsa gerçek çağrı, yoksa fallback."""
    api_key = os.environ.get('GEMINI_API_KEY')
    prompt = (
        "Aşağıdaki soruyu detaylı ve adım adım çöz. Türkçe açıkla.\n"
        "Soru Başlığı: " + (question.title or "") + "\n"
        "Soru İçeriği:\n" + (question.content or "") + "\n"
        "Format:\n- Yaklaşım\n- Adım adım çözüm\n- Varsa formüller\n- Sonuç\n"
    )
    if api_key:
        try:
            import google.generativeai as genai
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel("gemini-2.0-flash")
            resp = model.generate_content(prompt)
            text = getattr(resp, 'text', None)
            if text:
                return text
        except Exception:
            pass
    # Fallback
    base = "Yapay Zeka Çözüm Taslağı\n\n"
    if question.question_type == 'quiz':
        return base + (
            f"Soru Başlığı: {question.title}\n"
            "Yaklaşım: Şıkları eleme + tanım/formül doğrulama.\n"
            "Adımlar: Şık analizi → Uygun formül → Sonuç."
        )
    return base + (
        f"Soru Başlığı: {question.title}\n"
        "Çözüm Adımları: Varsayımlar → Tanımlar → Formüller → Ara Hesaplamalar → Sonuç.\n"
        "Kontrol: Birim, sınır koşulları ve mantık kontrolü yapın."
    )


def _split_choices(block: str):
    """A), B), C) ... veya A.,B.,C. şıklarını tespit eder, metni normalize eder."""
    # Şık başlıklarını satır başında ara
    parts = re.split(r"(?m)^(A\)|A\.|A\s*\-|A\:)\s*|^(B\)|B\.|B\s*\-|B\:)\s*|^(C\)|C\.|C\s*\-|C\:)\s*|^(D\)|D\.|D\s*\-|D\:)\s*|^(E\)|E\.|E\s*\-|E\:)\s*", block)
    # re.split dönen yapı etiketler de içerir; basit normalize için satırları birleştir
    cleaned = re.sub(r"\n{2,}", "\n", block).strip()
    return cleaned


def parse_questions_from_text(text: str):
    """Metinden birden çok soruyu yakala. 'Soru 1', '1.', '1)' vb. başlıklarına göre böler.
    Quiz ise şıkları korur. Basit heuristics; gerekirse geliştirilebilir.
    Dönen değer: list[ {title, content, type} ]
    """
    if not text:
        return []
    # Normalize line endings
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    # Soruları başlıklarına göre böl: "Soru 1", "1)", "1." satır başları
    chunks = re.split(r"(?m)^\s*(Soru\s*\d+\s*:?|\d+\)|\d+\.)\s*", text)
    # re.split ilk boş/giriş metni de verebilir; tekrar birleştirirken başlıkları ekle
    results = []
    buffer = []
    title_counter = 0
    for part in chunks:
        if not part:
            continue
        # Başlık gibi görünen parça: 'Soru 1', '1)' vb.
        if re.fullmatch(r"(Soru\s*\d+\s*:?|\d+\)|\d+\.)", part.strip(), flags=re.IGNORECASE):
            # önceki buffer'ı kaydet
            if buffer:
                body = "".join(buffer).strip()
                if body:
                    title_counter += 1
                    results.append({
                        'title': f'Soru {title_counter}',
                        'content': _split_choices(body),
                        'type': 'quiz' if re.search(r"(?m)^(A\)|A\.|B\)|B\.|C\)|C\.|D\)|D\.|E\)|E\.)", body) else 'solution'
                    })
            buffer = []
        else:
            buffer.append(part)
    if buffer:
        body = "".join(buffer).strip()
        if body:
            title_counter += 1
            results.append({
                'title': f'Soru {title_counter}',
                'content': _split_choices(body),
                'type': 'quiz' if re.search(r"(?m)^(A\)|A\.|B\)|B\.|C\)|C\.|D\)|D\.|E\)|E\.)", body) else 'solution'
            })
    return results


def parse_questions_from_docx(file_obj):
    """DOCX içeriğini okuyup text'e çevirir ve parse_questions_from_text ile ayrıştırır."""
    try:
        from docx import Document
        document = Document(file_obj)
        full_text = []
        for p in document.paragraphs:
            full_text.append(p.text)
        text = "\n".join(full_text)
        return parse_questions_from_text(text)
    except Exception:
        return []


# --------- QUIZ PARSING ---------
QUIZ_CHOICE_PATTERN = re.compile(r"^(?P<label>[ABCDE])\)|^(?P<label2>[ABCDE])\.|^(?P<label3>[ABCDE])\s*\-\s*", re.MULTILINE)


def _extract_text_from_pdf(file_obj) -> str:
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(file_obj)
        pages = []
        for p in reader.pages:
            pages.append(p.extract_text() or '')
        return "\n".join(pages)
    except Exception:
        return ''


def parse_quiz_from_text(text: str):
    """Soru ve şıkları çıkar, doğru şık yıldızlı (*A) gibi işaretlenmişse algılar."""
    items = parse_questions_from_text(text)
    quiz_items = []
    for it in items:
        content = it['content']
        # Şıkları bul
        blocks = re.split(r"(?m)^(A\)|A\.|B\)|B\.|C\)|C\.|D\)|D\.|E\)|E\.)\s*", content)
        # blocks karışık gelir; daha basit: satırlara bak
        lines = [l.strip() for l in content.splitlines() if l.strip()]
        question_lines = []
        choices = []
        for line in lines:
            m = re.match(r"^(\*?)([ABCDE])\)?\.?\s*(.+)$", line)
            if m:
                star, label, ctext = m.groups()
                choices.append({'label': label, 'text': ctext.strip(), 'correct': bool(star)})
            else:
                question_lines.append(line)
        question_text = " ".join(question_lines) or it['content']
        correct = None
        for c in choices:
            if c['correct']:
                correct = c['label']
                break
        quiz_items.append({'title': it['title'], 'question': question_text, 'choices': choices, 'correct': correct})
    return quiz_items


def parse_quiz_from_docx(file_obj):
    items = parse_questions_from_docx(file_obj)
    joined = "\n\n".join([f"{it['title']}\n{it['content']}" for it in items])
    return parse_quiz_from_text(joined)


def parse_quiz_from_pdf(file_obj):
    text = _extract_text_from_pdf(file_obj)
    return parse_quiz_from_text(text)


def create_quiz_from_items(course: Course, teacher, title: str, quiz_type: str, items: list, duration_minutes: int = 20) -> Quiz:
    quiz = Quiz.objects.create(course=course, created_by=teacher, title=title, quiz_type=quiz_type, duration_minutes=duration_minutes, is_published=False)
    order = 1
    for it in items:
        q = QuizQuestion.objects.create(quiz=quiz, order=order, text=it.get('question') or it.get('title'))
        correct_obj = None
        for ch in it.get('choices') or []:
            choice = QuizChoice.objects.create(question=q, label=ch.get('label',''), text=ch.get('text',''))
            if it.get('correct') and ch.get('label') == it['correct']:
                correct_obj = choice
        if correct_obj:
            q.correct_choice = correct_obj
            q.save(update_fields=['correct_choice'])
        order += 1
    return quiz


# --------- FEEDBACK GENERATION ---------
def generate_feedback(student, assignment, rubric: Dict[str, int] | None = None) -> str:
    """Template tabanlı basit geri bildirim üretimi.
    - templates/utils/feedback_templates/default.txt dosyasını kullanır.
    - rubric: {'içerik': 40, 'biçim': 20, 'özgünlük': 40} gibi dağılım.
    """
    rubric = rubric or {'İçerik': 40, 'Biçim': 20, 'Özgünlük': 40}
    # Derlem konteksti
    ctx = {
        'student_name': getattr(student, 'full_name', ''),
        'course_code': assignment.group.course.code,
        'course_name': assignment.group.course.name,
        'assignment_title': assignment.title,
        'due_date': assignment.due_date,
        'rubric': rubric,
    }
    # Şablon yolu
    try:
        from django.template.loader import render_to_string
        text = render_to_string('utils/feedback_templates/default.txt', ctx)
        return text.strip()
    except Exception:
        # Fallback düz metin
        lines = [
            f"{ctx['student_name']} için geri bildirim - {ctx['course_code']} {ctx['course_name']}",
            f"Ödev: {ctx['assignment_title']}",
            "",
            "Güçlü Yönler: Konu hakimiyeti ve örneklerle destekleme başarılı.",
            "Geliştirilecek Alanlar: Biçimlendirme, kaynakça ve görsel destek artırılabilir.",
            "Öneriler: Rubriğe göre eksik kalan kısımlara odaklanın.",
        ]
        return "\n".join(lines)


# --------- PLAGIARISM CHECK ---------
def _shingles(text: str, n: int = 5) -> set:
    s = re.sub(r"\s+", " ", (text or '').lower()).strip()
    return set(s[i:i+n] for i in range(max(0, len(s) - n + 1)))


def check_plagiarism(submission_text: str, corpus: List[str]) -> Dict[str, Any]:
    """Basit benzerlik ölçümleri: Jaccard(shingles) ve difflib ratio.
    Döner: {'max_similarity': float, 'items': [{'jaccard': x, 'ratio': y, 'index': i}]}
    """
    if not submission_text or not corpus:
        return {'max_similarity': 0.0, 'items': []}
    sub_sh = _shingles(submission_text)
    results = []
    max_sim = 0.0
    for idx, doc in enumerate(corpus):
        doc_sh = _shingles(doc)
        inter = len(sub_sh & doc_sh)
        union = len(sub_sh | doc_sh) or 1
        jacc = inter / union
        ratio = difflib.SequenceMatcher(None, submission_text, doc).ratio()
        sim = max(jacc, ratio)
        max_sim = max(max_sim, sim)
        results.append({'index': idx, 'jaccard': round(jacc, 4), 'ratio': round(ratio, 4)})
    return {'max_similarity': round(max_sim, 4), 'items': results}


# --------- FILE TEXT EXTRACTION (PDF/DOCX/TXT) ---------
def extract_text_from_upload(file_field) -> str:
    """Extract text from uploaded file robustly (PDF, DOCX, TXT fallback).
    - Tries PyPDF2 first for PDFs; if empty, tries pdfminer.six if available.
    - For DOCX, uses python-docx.
    - Otherwise attempts utf-8/latin-1 decode.
    """
    if not file_field or not hasattr(file_field, 'name'):
        return ''
    name = (file_field.name or '').lower()
    try:
        if name.endswith('.pdf'):
            # Try PyPDF2
            try:
                with file_field.open('rb') as f:
                    text = _extract_text_from_pdf(f)
                if text and text.strip():
                    return text
            except Exception:
                pass
            # Try pdfminer.six
            try:
                from pdfminer.high_level import extract_text as pdfminer_extract
                with file_field.open('rb') as f:
                    data = f.read()
                from io import BytesIO
                return pdfminer_extract(BytesIO(data)) or ''
            except Exception:
                return ''
        elif name.endswith('.docx'):
            try:
                from docx import Document
                with file_field.open('rb') as f:
                    data = f.read()
                from io import BytesIO
                doc = Document(BytesIO(data))
                return "\n".join([p.text or '' for p in doc.paragraphs])
            except Exception:
                return ''
        else:
            # Plain text or others
            try:
                with file_field.open('rb') as f:
                    raw = f.read()
                try:
                    return raw.decode('utf-8', errors='ignore')
                except Exception:
                    return raw.decode('latin-1', errors='ignore')
            except Exception:
                return ''
    except Exception:
        return ''


# ============================================================================
# UDEMY PLATFORM SERVICES - Online Kurs Sistemi İş Mantığı
# ============================================================================

class LessonProgressService:
    """Ders ilerleme takibi iş mantığı"""
    
    @staticmethod
    def update_video_progress(student, lesson, watched_duration):
        """
        Video izleme ilerlemesini güncelle
        %80 izlenmişse tamamlandı say
        """
        from .models import LessonProgress, CourseEnrollment
        
        enrollment = CourseEnrollment.objects.get(
            student=student, 
            course=lesson.module.course
        )
        
        progress, created = LessonProgress.objects.get_or_create(
            student=student,
            lesson=lesson,
            enrollment=enrollment
        )
        
        progress.watched_duration = watched_duration
        
        # Tamamlanma yüzdesini hesapla
        if lesson.video_duration > 0:
            progress.completion_percentage = min(
                (watched_duration / lesson.video_duration) * 100, 
                100
            )
        else:
            progress.completion_percentage = 0
        
        # %80 izlenmişse tamamlandı say
        if progress.completion_percentage >= 80:
            progress.status = 'completed'
            if not progress.completed_at:
                progress.completed_at = timezone.now()
        else:
            progress.status = 'in_progress'
        
        if not progress.started_at:
            progress.started_at = timezone.now()
        
        progress.save()
        
        # Kurs genel ilerlemesini güncelle
        LessonProgressService.update_course_progress(student, lesson.module.course)
        
        return progress
    
    @staticmethod
    def mark_pdf_completed(student, lesson):
        """PDF okundu olarak işaretle"""
        from .models import LessonProgress, CourseEnrollment
        
        enrollment = CourseEnrollment.objects.get(
            student=student, 
            course=lesson.module.course
        )
        
        progress, created = LessonProgress.objects.get_or_create(
            student=student,
            lesson=lesson,
            enrollment=enrollment
        )
        
        progress.status = 'completed'
        progress.completion_percentage = 100
        
        if not progress.completed_at:
            progress.completed_at = timezone.now()
        if not progress.started_at:
            progress.started_at = timezone.now()
        
        progress.save()
        
        LessonProgressService.update_course_progress(student, lesson.module.course)
        return progress
    
    @staticmethod
    def mark_text_completed(student, lesson):
        """Metin içerik okundu olarak işaretle"""
        return LessonProgressService.mark_pdf_completed(student, lesson)
    
    @staticmethod
    def update_quiz_progress(student, lesson, quiz_attempt):
        """
        Quiz sonucunu lesson progress'e kaydet
        Quiz geçilmişse tamamlandı say
        """
        from .models import LessonProgress, CourseEnrollment
        
        enrollment = CourseEnrollment.objects.get(
            student=student, 
            course=lesson.module.course
        )
        
        progress, created = LessonProgress.objects.get_or_create(
            student=student,
            lesson=lesson,
            enrollment=enrollment
        )
        
        progress.quiz_attempt = quiz_attempt
        progress.quiz_score = quiz_attempt.score or 0
        
        # Quiz geçme notu %60
        passing_score = 60
        progress.quiz_passed = progress.quiz_score >= passing_score
        
        if progress.quiz_passed:
            progress.status = 'completed'
            progress.completion_percentage = 100
            if not progress.completed_at:
                progress.completed_at = timezone.now()
        else:
            progress.status = 'in_progress'
            progress.completion_percentage = 50
        
        if not progress.started_at:
            progress.started_at = timezone.now()
        
        progress.save()
        
        if progress.quiz_passed:
            LessonProgressService.update_course_progress(student, lesson.module.course)
        
        return progress
    
    @staticmethod
    def submit_optional_assignment(student, lesson, file, notes=''):
        """
        Opsiyonel ödev gönder
        Ödev opsiyonel olduğu için not verilmez, sadece işaretlenir
        """
        from .models import LessonProgress, CourseEnrollment
        
        enrollment = CourseEnrollment.objects.get(
            student=student, 
            course=lesson.module.course
        )
        
        progress, created = LessonProgress.objects.get_or_create(
            student=student,
            lesson=lesson,
            enrollment=enrollment
        )
        
        progress.assignment_submitted = True
        progress.assignment_file = file
        progress.assignment_notes = notes
        progress.assignment_submitted_at = timezone.now()
        
        # Ödev gönderilince tamamlandı say
        progress.status = 'completed'
        progress.completion_percentage = 100
        
        if not progress.completed_at:
            progress.completed_at = timezone.now()
        if not progress.started_at:
            progress.started_at = timezone.now()
        
        progress.save()
        
        # Ödev opsiyonel olsa da kurs ilerlemesini etkiler
        LessonProgressService.update_course_progress(student, lesson.module.course)
        
        return progress
    
    @staticmethod
    def update_course_progress(student, course):
        """
        Kurs genel ilerlemesini güncelle
        Tüm zorunlu dersler tamamlandıysa sınav erişimi aç
        """
        from .models import LessonProgress, CourseEnrollment, Lesson
        from django.db.models import Q, Count
        
        enrollment = CourseEnrollment.objects.get(student=student, course=course)
        
        # Toplam zorunlu ders sayısı
        total_mandatory_lessons = Lesson.objects.filter(
            module__course=course,
            is_mandatory=True
        ).count()
        
        # Tamamlanan zorunlu dersler
        completed_mandatory_lessons = LessonProgress.objects.filter(
            student=student,
            lesson__module__course=course,
            lesson__is_mandatory=True,
            status='completed'
        ).count()
        
        enrollment.total_lessons_count = total_mandatory_lessons
        enrollment.completed_lessons_count = completed_mandatory_lessons
        
        # İlerleme yüzdesini hesapla
        if total_mandatory_lessons > 0:
            enrollment.progress_percentage = (
                completed_mandatory_lessons / total_mandatory_lessons
            ) * 100
        else:
            enrollment.progress_percentage = 0
        
        # Tüm zorunlu dersler tamamlandıysa sınav erişimi aç
        if enrollment.progress_percentage >= 100:
            if not enrollment.is_eligible_for_exam:
                enrollment.is_eligible_for_exam = True
                enrollment.exam_access_date = timezone.now()
        
        enrollment.save()
        return enrollment
    
    @staticmethod
    def get_next_lesson(student, course):
        """
        Öğrencinin bir sonraki tamamlaması gereken dersi bul
        """
        from .models import LessonProgress, Lesson, CourseModule
        
        # Son erişilen ders
        last_progress = LessonProgress.objects.filter(
            student=student,
            lesson__module__course=course
        ).order_by('-last_accessed').first()
        
        if not last_progress:
            # Hiç ders başlanmamışsa ilk modülün ilk dersi
            first_module = CourseModule.objects.filter(
                course=course,
                is_active=True
            ).order_by('order').first()
            
            if first_module:
                return first_module.lessons.filter(is_mandatory=True).order_by('order').first()
            return None
        
        # Aynı modülde sonraki ders
        next_lesson = Lesson.objects.filter(
            module=last_progress.lesson.module,
            order__gt=last_progress.lesson.order,
            is_mandatory=True
        ).order_by('order').first()
        
        if next_lesson:
            return next_lesson
        
        # Sonraki modülün ilk dersi
        next_module = CourseModule.objects.filter(
            course=course,
            order__gt=last_progress.lesson.module.order,
            is_active=True
        ).order_by('order').first()
        
        if next_module:
            return next_module.lessons.filter(is_mandatory=True).order_by('order').first()
        
        return None


class ExamService:
    """Sınav yönetimi iş mantığı"""
    
    @staticmethod
    def can_take_exam(student, course):
        """Öğrenci sınava girebilir mi kontrol et"""
        from .models import CourseEnrollment
        
        try:
            enrollment = CourseEnrollment.objects.get(
                student=student, 
                course=course
            )
            return enrollment.can_take_exam
        except CourseEnrollment.DoesNotExist:
            return False
    
    @staticmethod
    def get_remaining_attempts(student, course):
        """Kalan deneme hakkını hesapla"""
        from .models import CourseEnrollment, CourseExam, ExamAttempt
        
        try:
            enrollment = CourseEnrollment.objects.get(student=student, course=course)
            exam = course.final_exam
            
            used_attempts = ExamAttempt.objects.filter(
                enrollment=enrollment,
                exam=exam
            ).count()
            
            return max(0, exam.max_attempts - used_attempts)
        except Exception:
            return 0
    
    @staticmethod
    def start_exam(student, course):
        """
        Sınavı başlat
        - Deneme hakkı kontrolü
        - QuizAttempt ve ExamAttempt oluştur
        """
        from .models import (
            CourseEnrollment, CourseExam, ExamAttempt, 
            QuizAttempt
        )
        from django.core.exceptions import ValidationError
        
        # Enrollment kontrolü
        try:
            enrollment = CourseEnrollment.objects.get(student=student, course=course)
        except CourseEnrollment.DoesNotExist:
            raise ValidationError('Bu kursa kayıtlı değilsiniz')
        
        # Sınava uygun mu?
        if not enrollment.can_take_exam:
            raise ValidationError('Henüz sınava girebilmek için tüm içeriği tamamlamadınız')
        
        # Exam kontrolü
        try:
            exam = course.final_exam
        except Exception:
            raise ValidationError('Bu kursun sınavı bulunmuyor')
        
        # Deneme hakkı kontrolü
        attempts_count = ExamAttempt.objects.filter(
            enrollment=enrollment,
            exam=exam
        ).count()
        
        if attempts_count >= exam.max_attempts:
            raise ValidationError(f'Maksimum deneme sayısına ({exam.max_attempts}) ulaştınız')
        
        # QuizAttempt oluştur
        quiz_attempt = QuizAttempt.objects.create(
            quiz=exam.quiz,
            student=student
        )
        
        # ExamAttempt oluştur
        exam_attempt = ExamAttempt.objects.create(
            enrollment=enrollment,
            exam=exam,
            student=student,
            attempt_number=attempts_count + 1,
            quiz_attempt=quiz_attempt
        )
        
        return exam_attempt
    
    @staticmethod
    def complete_exam(exam_attempt):
        """
        Sınavı tamamla ve değerlendir
        - Cevapları kontrol et
        - Geçti/kaldı durumunu belirle
        - Geçtiyse sertifika oluştur
        """
        from .models import QuizAnswer
        
        quiz_attempt = exam_attempt.quiz_attempt
        quiz_attempt.is_submitted = True
        quiz_attempt.completed_at = timezone.now()
        
        # Cevapları değerlendir
        total_questions = quiz_attempt.quiz.questions.count()
        
        if total_questions == 0:
            raise ValueError('Sınavda soru bulunmuyor')
        
        # Doğru cevapları say
        correct_answers = QuizAnswer.objects.filter(
            attempt=quiz_attempt,
            is_correct=True
        ).count()
        
        # Puanı hesapla
        score = (correct_answers / total_questions) * 100
        quiz_attempt.score = score
        quiz_attempt.save()
        
        exam_attempt.score = score
        exam_attempt.completed_at = timezone.now()
        
        # Geçti mi?
        if score >= exam_attempt.exam.passing_score:
            exam_attempt.status = 'passed'
            
            # Enrollment'ı tamamla
            enrollment = exam_attempt.enrollment
            enrollment.status = 'completed'
            enrollment.completed_at = timezone.now()
            enrollment.save()
            
            # Sertifika oluştur
            CertificateService.generate_certificate(exam_attempt)
        else:
            exam_attempt.status = 'failed'
        
        exam_attempt.save()
        return exam_attempt


class CertificateService:
    """Sertifika oluşturma ve yönetimi"""
    
    @staticmethod
    def generate_certificate(exam_attempt):
        """
        Otomatik PDF sertifika oluştur
        - Benzersiz sertifika ID
        - PDF dosyası oluştur (ReportLab ile)
        - Doğrulama URL'i
        """
        from .models import Certificate
        from django.conf import settings
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        import uuid
        import os
        
        enrollment = exam_attempt.enrollment
        student = enrollment.student
        course = enrollment.course
        
        # Benzersiz sertifika ID oluştur
        cert_id = f"CERT-{course.code}-{student.school_number}-{uuid.uuid4().hex[:8].upper()}"
        
        # Certificate kaydı oluştur
        certificate = Certificate.objects.create(
            enrollment=enrollment,
            student=student,
            course=course,
            certificate_id=cert_id,
            exam_score=exam_attempt.score,
            completion_date=exam_attempt.completed_at,
            verification_url=f"{settings.SITE_URL}/verify-certificate/{cert_id}/"
        )
        
        # PDF dosyası oluştur
        pdf_filename = f'{cert_id}.pdf'
        pdf_path = os.path.join('certificates', pdf_filename)
        full_path = os.path.join(settings.MEDIA_ROOT, pdf_path)
        
        # Dizini oluştur
        os.makedirs(os.path.dirname(full_path), exist_ok=True)
        
        # Canvas oluştur (landscape A4)
        c = canvas.Canvas(full_path, pagesize=landscape(A4))
        width, height = landscape(A4)
        
        # Arka plan rengi
        c.setFillColorRGB(0.95, 0.95, 0.98)
        c.rect(0, 0, width, height, fill=1, stroke=0)
        
        # Çerçeve
        c.setStrokeColorRGB(0.2, 0.4, 0.8)
        c.setLineWidth(3)
        c.rect(2*cm, 2*cm, width-4*cm, height-4*cm, fill=0, stroke=1)
        
        # Başlık
        c.setFillColorRGB(0.1, 0.2, 0.6)
        c.setFont("Helvetica-Bold", 48)
        c.drawCentredString(width/2, height - 5*cm, "SERTİFİKA")
        
        # Alt başlık
        c.setFillColorRGB(0, 0, 0)
        c.setFont("Helvetica", 18)
        c.drawCentredString(width/2, height - 7*cm, "Bu belge ile")
        
        # Öğrenci adı
        c.setFont("Helvetica-Bold", 32)
        c.setFillColorRGB(0.1, 0.3, 0.7)
        c.drawCentredString(width/2, height - 9.5*cm, student.full_name)
        
        # Kurs bilgisi
        c.setFillColorRGB(0, 0, 0)
        c.setFont("Helvetica", 20)
        c.drawCentredString(width/2, height - 11.5*cm, course.name)
        c.setFont("Helvetica", 16)
        c.drawCentredString(width/2, height - 12.8*cm, f"({course.code})")
        
        # Açıklama
        c.setFont("Helvetica", 16)
        c.drawCentredString(width/2, height - 14.5*cm, "kursunu başarıyla tamamladığını onaylar.")
        
        # Detaylar
        c.setFont("Helvetica", 13)
        c.drawCentredString(width/2, height - 16.5*cm, 
                          f"Sınav Puanı: {exam_attempt.score:.1f}/100")
        c.drawCentredString(width/2, height - 17.5*cm, 
                          f"Tamamlanma Tarihi: {enrollment.completed_at.strftime('%d.%m.%Y')}")
        
        # Sertifika numarası ve doğrulama
        c.setFont("Helvetica", 10)
        c.setFillColorRGB(0.4, 0.4, 0.4)
        c.drawCentredString(width/2, height - 19*cm, 
                          f"Sertifika No: {cert_id}")
        c.drawCentredString(width/2, height - 19.7*cm, 
                          f"Doğrulama: {certificate.verification_url}")
        
        # Kaydet
        c.save()
        
        # Certificate'e PDF dosyasını ekle
        certificate.certificate_file = pdf_path
        certificate.save()
        
        # Enrollment'ı güncelle
        enrollment.certificate_issued = True
        enrollment.certificate_issued_at = timezone.now()
        enrollment.save()
        
        return certificate
    
    @staticmethod
    def verify_certificate(certificate_id):
        """Sertifika doğrula"""
        from .models import Certificate
        
        try:
            certificate = Certificate.objects.get(
                certificate_id=certificate_id,
                is_valid=True
            )
            return {
                'valid': True,
                'student': certificate.student.full_name,
                'course': certificate.course.name,
                'issue_date': certificate.issue_date,
                'score': certificate.exam_score
            }
        except Certificate.DoesNotExist:
            return {'valid': False, 'message': 'Sertifika bulunamadı'}
    
    @staticmethod
    def revoke_certificate(certificate, reason=''):
        """Sertifikayı iptal et"""
        certificate.is_valid = False
        certificate.revoked_at = timezone.now()
        certificate.revoked_reason = reason
        certificate.save()
        return certificate


class CourseEnrollmentService:
    """Kurs kayıt işlemleri"""
    
    @staticmethod
    def enroll_student(student, course):
        """
        Öğrenciyi kursa kaydet
        - Zaten kayıtlı mı kontrol et
        - Toplam ders sayısını hesapla
        """
        from .models import CourseEnrollment, Lesson
        from django.core.exceptions import ValidationError
        
        # Zaten kayıtlı mı?
        if CourseEnrollment.objects.filter(student=student, course=course).exists():
            raise ValidationError('Bu kursa zaten kayıtlısınız')
        
        # Kurs aktif mi?
        if course.status != 'active':
            raise ValidationError('Bu kurs şu anda aktif değil')
        
        # Online kurs mu?
        if course.course_type != 'online':
            raise ValidationError('Sadece online kurslara kaydolabilirsiniz')
        
        # Toplam zorunlu ders sayısı
        total_lessons = Lesson.objects.filter(
            module__course=course,
            is_mandatory=True
        ).count()
        
        # Kayıt oluştur
        enrollment = CourseEnrollment.objects.create(
            student=student,
            course=course,
            total_lessons_count=total_lessons
        )
        
        return enrollment
    
    @staticmethod
    def get_student_dashboard(student):
        """
        Öğrenci dashboard verilerini hazırla
        - Aktif kurslar
        - İlerleme durumları
        - Sertifikalar
        """
        from .models import CourseEnrollment
        from django.db.models import Prefetch
        
        # Aktif kurslar
        active_enrollments = CourseEnrollment.objects.filter(
            student=student,
            status='active'
        ).select_related('course').prefetch_related(
            'lesson_progress',
            'exam_attempts'
        )
        
        dashboard_data = []
        for enrollment in active_enrollments:
            # Son erişilen ders
            last_lesson_progress = enrollment.lesson_progress.order_by(
                '-last_accessed'
            ).first()
            
            # Sonraki ders
            next_lesson = LessonProgressService.get_next_lesson(
                student, 
                enrollment.course
            )
            
            dashboard_data.append({
                'enrollment': enrollment,
                'last_lesson': last_lesson_progress.lesson if last_lesson_progress else None,
                'next_lesson': next_lesson,
                'exam_attempts_count': enrollment.exam_attempts.count(),
            })
        
        # Tamamlanan kurslar
        completed_enrollments = CourseEnrollment.objects.filter(
            student=student,
            status='completed',
            certificate_issued=True
        ).select_related('course', 'certificate')
        
        return {
            'active_courses': dashboard_data,
            'completed_courses': completed_enrollments,
            'total_active': active_enrollments.count(),
            'total_completed': completed_enrollments.count(),
        }